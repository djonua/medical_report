import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
import json

def create_medical_report(patient_name, report_text):
    """
    Создает медицинское заключение в формате DOCX
    
    Args:
        patient_name (str): ФИО пациента
        report_text (str): Текст заключения от ChatGPT
    """
    # Создаем новый документ
    doc = Document()
    
    # Настраиваем стиль для заголовка
    title = doc.add_heading('Медицинское заключение', 0)
    title.alignment = 1  # Центрирование
    
    # Добавляем информацию о пациенте
    doc.add_paragraph(f'Пациент: {patient_name}')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph('')  # Пустая строка для разделения
    
    # Добавляем основной текст заключения
    doc.add_paragraph(report_text)
    
    # Создаем папку Reports, если она не существует
    if not os.path.exists('Reports'):
        os.makedirs('Reports')
    
    # Формируем имя файла (ФИО_дата.docx)
    safe_name = patient_name.replace(' ', '_')
    filename = f'Reports/{safe_name}_{datetime.now().strftime("%d_%m_%Y")}.docx'
    
    # Сохраняем документ
    doc.save(filename)
    return filename

def generate_report(patient_name, report_text):
    """
    Создает отчет на основе имени пациента и текста заключения
    
    Args:
        patient_name (str): ФИО пациента
        report_text (str): Текст заключения
    """
    try:
        report_file = create_medical_report(
            patient_name=patient_name,
            report_text=report_text
        )
        return report_file
    except Exception as e:
        print(f"Ошибка при создании отчета: {str(e)}")
        return None 